from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Open the love letter template
doc = Document('templates/love_letter_template.docx')

# Access the first paragraph (assuming placeholder is in the first paragraph)
paragraph = doc.paragraphs[0]

# Replace the placeholder with actual recipient's name and add an emoji
recipient_name = "My Love ❤️"
paragraph.clear()  # Clear the paragraph before adding new content
paragraph.add_run(f"🌹 My Dearest {recipient_name}, ❤️")

# Add the main content of the love letter
love_content = (
    "I wanted to take a moment to express the depth of my love for you. "
    "Your presence in my life brings me immeasurable joy and happiness. "
    "Every moment we share together is a treasure, and I cherish our love more than words can convey. ❤️"
)
doc.add_paragraph(love_content)

# Save the modified document as the final love letter
doc.save('final_love_letter.docx')
